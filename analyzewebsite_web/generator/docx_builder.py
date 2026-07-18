"""Render the analysis into a TEN Capital Website Analysis .docx.

Section order follows website_analysis_template.md exactly. The footer follows
the TEN Capital standard in CLAUDE.md: Open Sans 7pt, centered, single line,
with the page number and the logo inline.
"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

BODY_FONT = "Open Sans"
FOOTER_PT = 7
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x5A, 0x64, 0x72)

ASSETS = Path(__file__).resolve().parent.parent / "assets"
LOGO = ASSETS / "TEN_Capital_logo_footer.png"

SCORECARD_ROWS = [
    ("scientific_technical_credibility", "Scientific / Technical Credibility"),
    ("founder_credibility", "Founder Credibility"),
    ("product_positioning", "Product Positioning"),
    ("market_opportunity_communication", "Market Opportunity Communication"),
    ("commercialization_story", "Commercialization Story"),
    ("traction_evidence", "Traction Evidence"),
    ("investor_readiness", "Investor Readiness"),
    ("fundraising_supportiveness", "Fundraising Supportiveness"),
]

PROBE_ROWS = [
    ("why_now", "Why Now?"),
    ("why_this_team", "Why This Team?"),
    ("why_this_market", "Why This Market?"),
    ("why_this_product_wins", "Why This Product Wins?"),
    ("why_this_becomes_large", "Why This Becomes Large?"),
]

DISCLAIMER = (
    "Information presented is for informational purposes only and does not "
    "constitute an offer to sell or solicitation of an offer to buy securities."
)


class DocumentError(Exception):
    """Raised when the .docx cannot be produced."""


def build_analysis_docx(data: dict, analysis_date: str) -> bytes:
    """Return the finished .docx as bytes."""
    try:
        doc = Document()
        _configure_base_style(doc)
        _configure_page(doc)

        company = (data.get("company_name") or "Company").strip()
        title = f"{company} - Website Analysis"
        _build_header(doc.sections[0])
        _build_footer(doc.sections[0], title, analysis_date)

        _title_block(doc, title, data)
        _executive_summary(doc, data)
        _whats_working(doc, data)
        _gaps(doc, data)
        _improvements(doc, data)
        _financial_storytelling(doc, data)
        _reg_d(doc, data)
        _ux(doc, data)
        _overall(doc, data)
        _provenance(doc, data)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except DocumentError:
        raise
    except Exception as exc:  # python-docx raises broadly on malformed input
        raise DocumentError(f"Could not build the document: {exc}") from exc


# --------------------------------------------------------------------- chrome


def _configure_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    # python-docx sets only the latin font; set the east-asian slot too or Word
    # silently substitutes for any non-ASCII character.
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, bold in (
        ("Heading 1", 15, ACCENT, True),
        ("Heading 2", 12, ACCENT, True),
        ("Heading 3", 10.5, RGBColor(0x12, 0x16, 0x1C), True),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(5)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)


def _configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _build_header(section) -> None:
    """Empty header with the tab stops specified in CLAUDE.md."""
    header = section.header
    header.is_linked_to_previous = False
    paragraphs = header.paragraphs or [header.add_paragraph()]
    while len(paragraphs) < 2:
        paragraphs.append(header.add_paragraph())
    for para in paragraphs[:2]:
        para.text = ""
        stops = para.paragraph_format.tab_stops
        stops.add_tab_stop(Twips(4677), WD_TAB_ALIGNMENT.CENTER)
        stops.add_tab_stop(Twips(9355), WD_TAB_ALIGNMENT.RIGHT)


def _build_footer(section, title: str, analysis_date: str) -> None:
    """TEN Capital footer: title, page number, compiled-on line, logo."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ""
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _footer_run(para, f" {title}                     ")
    _page_number_field(para)
    _footer_run(para, f"                  Compiled on {analysis_date} by TEN Capital Network    ")

    if LOGO.exists():
        run = para.add_run()
        run.add_picture(str(LOGO), width=Inches(0.67), height=Inches(0.25))
    # A missing logo degrades the footer but must not fail the download.


def _footer_run(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(FOOTER_PT)
    return run


def _page_number_field(paragraph) -> None:
    """Insert a live { PAGE } field — python-docx has no API for this."""
    run = paragraph.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(FOOTER_PT)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


# --------------------------------------------------------------------- helpers


def _para(doc, text: str, *, italic=False, bold=False, size=None, color=None, after=None):
    para = doc.add_paragraph()
    run = para.add_run(text or "")
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    return para


def _bullets(doc, items) -> None:
    for item in items or []:
        if str(item).strip():
            doc.add_paragraph(str(item).strip(), style="List Bullet")


def _table(doc, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            run.font.size = Pt(9.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return table


# --------------------------------------------------------------------- sections


def _title_block(doc, title: str, data: dict) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    meta = [data.get("website_url") or "", data.get("sector") or ""]
    _para(doc, "  •  ".join(m for m in meta if m), size=9, color=MUTED, after=12)


def _executive_summary(doc, data: dict) -> None:
    summary = data.get("executive_summary") or {}
    doc.add_heading("Executive Summary", level=1)
    _para(doc, summary.get("what_the_site_does_well", ""))
    _para(doc, summary.get("who_the_site_serves", ""))

    score = summary.get("readiness_score", 0)
    score_text = f"{score:.1f}".rstrip("0").rstrip(".") if isinstance(score, float) else str(score)
    _para(doc, f"Current Investor Readiness Score: {score_text}/10", bold=True, after=10)

    doc.add_heading("Category Scorecard", level=2)
    scorecard = data.get("scorecard") or {}
    _table(
        doc,
        ["Category", "Assessment"],
        [[label, scorecard.get(key, "—")] for key, label in SCORECARD_ROWS],
        widths=[4.0, 2.5],
    )


def _whats_working(doc, data: dict) -> None:
    items = data.get("whats_working") or []
    if not items:
        return
    doc.add_heading("What's Working", level=1)
    for i, item in enumerate(items, start=1):
        doc.add_heading(f"{i}. {item.get('title', '').strip()}", level=2)
        _para(doc, item.get("body", ""))
        _bullets(doc, item.get("supporting_points"))

        callout = item.get("callout") or {}
        label = (callout.get("label") or "None").strip()
        if label and label != "None":
            doc.add_heading(label, level=3)
            if callout.get("lead_in"):
                _para(doc, callout["lead_in"])
            _bullets(doc, callout.get("points"))
            if callout.get("close"):
                _para(doc, callout["close"])


def _gaps(doc, data: dict) -> None:
    gaps = data.get("gaps") or []
    if not gaps:
        return
    doc.add_heading("Gaps and Weaknesses", level=1)
    for i, gap in enumerate(gaps, start=1):
        doc.add_heading(f"{i}. {gap.get('title', '').strip()}", level=2)
        framing = gap.get("framing", "")
        if gap.get("is_largest_concern"):
            para = doc.add_paragraph()
            run = para.add_run("This is the largest investor concern. ")
            run.bold = True
            para.add_run(framing)
        else:
            _para(doc, framing)

        if gap.get("missing_items"):
            _para(doc, "The website does not clearly communicate:", after=2)
            _bullets(doc, gap["missing_items"])

        question = (gap.get("investor_question") or "").strip()
        if question:
            doc.add_heading("Investor Question", level=3)
            _para(doc, f"“{question.strip(chr(34))}”", italic=True)
        if gap.get("resolution_note"):
            _para(doc, gap["resolution_note"])

    probes = data.get("narrative_probes") or {}
    if probes:
        doc.add_heading("Investor-Facing Narrative", level=2)
        _para(
            doc,
            "How well the site answers the five questions an investor brings to it:",
            after=4,
        )
        _table(
            doc,
            ["Question", "Coverage"],
            [[label, probes.get(key, "—")] for key, label in PROBE_ROWS],
            widths=[4.0, 2.5],
        )


def _improvements(doc, data: dict) -> None:
    items = data.get("improvements") or []
    doc.add_heading("Actionable Improvements", level=1)
    for i, item in enumerate(items, start=1):
        doc.add_heading(f"Priority #{i}: {item.get('title', '').strip()}", level=2)
        _para(doc, item.get("intro", ""))
        for section in item.get("suggested_sections") or []:
            _para(doc, section.get("name", ""), bold=True, after=2)
            _bullets(doc, section.get("items"))
        example = item.get("example_block") or []
        if example:
            _para(doc, "Example:", italic=True, after=2)
            _bullets(doc, example)

    proof = data.get("commercial_proof_points") or []
    if proof:
        doc.add_heading("Commercial Proof Points", level=2)
        _para(
            doc,
            "Metrics this company should publish. Formats are placeholders, not "
            "reported figures.",
            after=4,
        )
        _table(
            doc,
            ["Metric", "Example Format"],
            [[p.get("metric", ""), p.get("example_format", "")] for p in proof],
            widths=[4.0, 2.5],
        )
        _para(doc, "Even modest numbers are better than no numbers.", italic=True)


def _financial_storytelling(doc, data: dict) -> None:
    fin = data.get("financial_storytelling") or {}
    doc.add_heading("Financial Storytelling (Without Violating Securities Rules)", level=1)
    _para(
        doc,
        fin.get("framing")
        or "The website can communicate growth without discussing a securities offering.",
    )
    if fin.get("commercial_milestones"):
        _para(doc, "Commercial Milestones", bold=True, after=2)
        _bullets(doc, fin["commercial_milestones"])
    if fin.get("market_opportunity_points"):
        _para(doc, "Market Opportunity — present:", bold=True, after=2)
        _bullets(doc, fin["market_opportunity_points"])
    _para(doc, "Avoid:", bold=True, after=2)
    _bullets(
        doc,
        [
            "Promises of investment returns",
            "Forecasted investor gains",
            "Promotional fundraising language",
        ],
    )


def _reg_d(doc, data: dict) -> None:
    reg = data.get("reg_d") or {}
    rule = (reg.get("applicable_rule") or "n/a").strip()
    doc.add_heading("SEC Regulation D Considerations", level=1)
    _para(doc, "(Assuming future U.S. fundraising activity.)", italic=True)
    if reg.get("jurisdiction_note"):
        _para(doc, reg["jurisdiction_note"])
    _para(doc, f"The site's current posture reads as: {rule}.", bold=True, after=8)

    if rule in ("506(b)", "both", "n/a"):
        doc.add_heading("If Using Rule 506(b)", level=2)
        _para(doc, "Avoid general solicitation. Recommended approach:", after=2)
        _bullets(
            doc,
            [
                "Public website remains informational",
                "No active fundraising language",
                "No investment opportunity pages",
                "Investor materials gated behind password access",
                "Access only after establishing substantive investor relationships",
            ],
        )
        _para(doc, "Appropriate structure", bold=True, after=4)
        _table(
            doc,
            ["Public Site", "Private Investor Portal"],
            [
                ["Company information", "Financials"],
                ["Product information", "Data room"],
                ["Team", "Pitch deck"],
                ["News", "Fundraising materials"],
            ],
            widths=[3.25, 3.25],
        )

    if rule in ("506(c)", "both", "n/a"):
        doc.add_heading("If Using Rule 506(c)", level=2)
        _para(doc, "Public discussion is permissible, but:", after=2)
        _bullets(
            doc,
            [
                "Avoid performance claims",
                "Avoid investment return projections",
                "Use clear securities disclaimers",
                "Verify accredited investor status before accepting investments",
            ],
        )
        _para(doc, "Recommended footer language", bold=True, after=2)
        _para(doc, DISCLAIMER, italic=True)

    _para(
        doc,
        "This section is general guidance, not legal advice. Confirm all securities "
        "matters with qualified counsel.",
        italic=True,
        size=9,
        color=MUTED,
    )


def _ux(doc, data: dict) -> None:
    ux = data.get("ux_enhancements") or {}
    doc.add_heading("Presentation & UX Enhancements", level=1)

    fold = ux.get("above_the_fold") or {}
    doc.add_heading("Homepage — Above the Fold", level=2)
    if fold.get("current"):
        _para(doc, f"Current: {fold['current']}")
    _para(doc, "Recommended:", bold=True, after=2)
    _bullets(
        doc,
        [
            "One-sentence value proposition",
            "Product image",
            "Key metrics",
            "Partner logos",
        ],
    )
    if fold.get("proposed_headline"):
        _para(doc, "Example:", italic=True, after=2)
        _para(doc, fold["proposed_headline"], bold=True, after=2)
        _bullets(doc, fold.get("sub_claims"))

    if ux.get("visual_hierarchy_observation"):
        doc.add_heading("Stronger Visual Hierarchy", level=2)
        _para(doc, f"Current: {ux['visual_hierarchy_observation']}")
        _para(doc, "Add:", bold=True, after=2)
        _bullets(
            doc,
            [
                "Infographics",
                "Timeline of milestones",
                "Product workflow diagrams",
                "Validation visuals",
            ],
        )

    if ux.get("leadership_section_observation"):
        doc.add_heading("Improve Leadership Section", level=2)
        _para(doc, ux["leadership_section_observation"])
        _para(doc, "Add:", bold=True, after=2)
        _bullets(
            doc,
            [
                "Professional headshots",
                "Founder achievements",
                "Publications",
                "Prior exits",
                "Academic or industry affiliations",
            ],
        )
        _para(doc, "Investors back people first.", italic=True)

    timeline = ux.get("milestone_timeline") or []
    if timeline:
        doc.add_heading("Add Milestone Timeline", level=2)
        _para(
            doc,
            "Dates below are drawn from the site. Confirm and extend before publishing.",
            after=4,
        )
        _table(
            doc,
            ["Year", "Milestone"],
            [[row.get("year", ""), row.get("milestone", "")] for row in timeline],
            widths=[1.2, 5.3],
        )
        _para(doc, "Investors love progress visualization.", italic=True)

    if ux.get("recommended_ctas"):
        doc.add_heading("Improve Calls-to-Action", level=2)
        if ux.get("cta_observation"):
            _para(doc, f"Current: {ux['cta_observation']}")
        _para(doc, "Add:", bold=True, after=2)
        _bullets(doc, ux["recommended_ctas"])


def _overall(doc, data: dict) -> None:
    overall = data.get("overall_assessment") or {}
    doc.add_heading("Overall Assessment", level=1)
    _para(doc, overall.get("current_position", ""))
    _para(doc, overall.get("primary_weakness", ""))

    fixes = overall.get("required_fixes") or []
    company = data.get("company_name") or "the company"
    if fixes:
        _para(doc, f"If {company} adds:", after=2)
        _bullets(doc, fixes)

    summary = data.get("executive_summary") or {}
    score = summary.get("readiness_score", 0)
    score_text = f"{score:.1f}".rstrip("0").rstrip(".") if isinstance(score, float) else str(score)
    target = overall.get("target_score_range", "")
    investors = overall.get("target_investor_types") or []
    audience = _join(investors)
    closing = (
        f"the website could move from a {score_text}/10 investor-readiness score to "
        f"approximately {target}/10"
    )
    if audience:
        closing += f", substantially improving first impressions with {audience}"
    _para(doc, closing + ".")


def _provenance(doc, data: dict) -> None:
    pages = data.get("pages_reviewed") or []
    if not pages:
        return
    doc.add_heading("Pages Reviewed", level=2)
    _para(
        doc,
        "This analysis is based on the following pages. Content behind logins, "
        "in downloadable files, or rendered only in the browser was not reviewed.",
        size=9,
        color=MUTED,
        after=4,
    )
    for url in pages:
        para = doc.add_paragraph(url, style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = MUTED


def _join(items: list[str]) -> str:
    items = [str(i).strip() for i in items if str(i).strip()]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    return ", ".join(items[:-1]) + f", and {items[-1]}"
